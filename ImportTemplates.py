
from docx import  Document
import re,os
SYMBOLS = {'{': '}', '[': ']'}
SYMBOLS_L, SYMBOLS_R = SYMBOLS.keys(), SYMBOLS.values()

class Templates():

    def __init__(self,filename):
        self.filename=filename #命令行模板文档
        self.completer=[]   #存放关键字补全列表
        self.keytoparameter={} # 命令字参数串对
        self.keytoparameterexplain={} #存放命令中参数说明 用于提示
        self.keywithregp={}
        self.ParameterToType={} #参数，类型描述对
        self.keytohelper={}  #保存（命令，命令文档）键值对
        self.TypeToReg={'int':'-?\d*.','ip_10':'\d+\.\d+\.\d+\.\d+','string':'[A-za-z0-9]+'}
        self.ParameterToReg={}
        self.REG=''
        self.imporT()



    def imporT(self):     #读取docx文档，将命令信息导入
        document=Document(self.filename);
        now='_'
        nowtitle=''
        temptokeepkey=[]
        #print(len(document.paragraphs))
        p=0   #word当前行
        helpS=0
        while p<len(document.paragraphs):   #遍历文档
           #print("while 1",p)

           """for r in p.runs:
               if r.bold:
                   self.completer.append(r.text)   """
           if document.paragraphs[p].style.name=='Heading 3':   #识别三级标题，导入单个命令
                nowtitle=document.paragraphs[p].text
                helpS=p  #helpS 记录当前命令word开始行，下面helpE记录结束行，使用-h命令时，显示开始行至结束行内容
                p=p+1
                while  p<len(document.paragraphs)and(document.paragraphs[p].style.name=="Normal"): #单个命令页面正文段
                    #print("while 2",p)

                    if "命令格式"in document.paragraphs[p].text:    #从命令格式一栏从提取命令关键字
                        p=p+1
                        while "参数说明" not in document.paragraphs[p].text: # “命令格式”至“参数说明”循环
                            #print("while 3 ",p)
                            count=0
                            k=''
                            Pstring=''
                            for r in document.paragraphs[p].runs:
                                #print("test issue "+r.text)
                                if r.bold:
                                    k=k+r.text
                                    self.completer.append(r.text)
                                    temptokeepkey.append(r.text)
                                    endpos=len(r.text)
                            #print(count)
                            pstring=document.paragraphs[p].text[endpos+1:]
                            self.keytoparameter[k]=pstring
                            #print(pstring)
                            p=p+1
                    if "参数说明"in document.paragraphs[p].text:
                        p = p + 1
                        explain=[]   #存放参数说明
                        while "视图" not in document.paragraphs[p].text:
                            for r in document.paragraphs[p].runs:
                                if r.italic:
                                    #print(r.text)
                                    endpos = len(r.text)
                                    tempp=document.paragraphs[p].text
                                    if "点分十进制" in tempp :           #存储参数的类型
                                        self.ParameterToType[r.text]="ip_10"
                                    elif "整数" in tempp:
                                        self.ParameterToType[r.text] = "int"
                                    elif "字符串" in tempp:
                                        self.ParameterToType[r.text] = "string"
                                    self.ParameterToReg[r.text]=self.TypeToReg[self.ParameterToType[r.text]]
                                    explain.append(tempp)

                                    break
                            p = p + 1
                        for key in temptokeepkey:
                            isfinsh=self.transform(self.keytoparameter[key])
                            if(isfinsh):
                             self.keywithregp[key]=self.REG
                            else:
                                print(nowtitle+" 参数字段不符规范")
                                print(key+" "+self.keytoparameter[key])
                            #print(self.REG)

                        #print(explain)
                    p=p+1
                helpE=p
                for k in temptokeepkey:                                            #保存命令帮助界面
                    self.keytohelper[k]=str(helpS)+"_"+str(helpE)
                    self.keytoparameterexplain[k]=explain.copy()
                temptokeepkey.clear()
                explain.clear()
                p=p-1
           p=p+1


    def analyze(self,text):
        text.strip()
        text=' '.join(text.split())
        key=''
        keyEndPos=0
        findit=False
        self.REG = ''
        for c in range(len(text)):       #查询是否有该命令
             if text[c]!=' ':
               key=key+text[c]
             if text[c]==' ':
                 if key in self.completer:
                     findit=True
                     keyEndPos=c
                     break
                 key=key+' '
             if c==(len(text)-1):
                if key in self.completer:
                    keyEndPos=c
                    findit=True
                    break
        if not findit:
            print(text+": command not found")
        else:
            if '-h' in text:
                gs,ge=self.keytohelper[key].split("_")
                document=Document(self.filename);
                for c in range(int(gs),int(ge)):
                    if document.paragraphs[c].text:
                        print(document.paragraphs[c].text)
            else:
                parameterstring=text[keyEndPos+1:]
                parameterstringreg=self.keywithregp[key]
                #print("参数串正则表达式：")
                #print(parameterstringreg)
                pattern = re.compile(r''+parameterstringreg)
                m=pattern.findall(parameterstring)
                #print(parameterstring)
                #print("匹配结果 "+str(m))
                if(m):    #匹配成功
                    try:
                        result=os.system(key+".exe "+parameterstring)
                    except:
                        print(key+" 命令没有执行程序")
                else:
                    print("参数输入格式不正确 请参照以下格式")
                    print(key+" "+self.keytoparameter[key])
                    for e in self.keytoparameterexplain[key]:
                        print(e)

    def transform(self,s):
        expect = []   #用作栈，存放期望匹配的括号
        arr = []       #用作栈，存放处理的参数串
        pos = []       #用作栈，存放括号的位置
        pnum = 1  # 标记一对括号对里参数的个数
        length = len(s)
        i = 0
        frontisor = False  # 判断此参数前面是否还有同等地位参数
        while i < length:
            c = s[i]
            if (len(pos) > 0):
                nowpos = pos[-1]
            if c in SYMBOLS_L:     #遇到左括号
                arr.append(c)
                expect.append(SYMBOLS[c])
                pos.append(len(arr) - 1)
            elif c in SYMBOLS_R:   #遇到右括号
                frontisor = False
                if len(expect)!=0 and c == expect[-1]:

                    if i + 1 < length:  # 判断是否带有*标志
                        if s[i + 1] == '*':
                            if c == ']':
                                temp = ''
                                for j in range(nowpos + 1, len(arr)):
                                    temp = temp + arr[j]
                                temp = "(" + temp + "){0," + str(pnum) + "}"
                                for j in range(nowpos, len(arr)):
                                    arr.pop()
                                arr.append(temp);
                                pnum = 1
                            elif c == '}':
                                temp = ''
                                for j in range(nowpos + 1, len(arr)):
                                    temp = temp + arr[j]
                                temp = "(" + temp + "){1," + str(pnum) + "}"
                                for j in range(nowpos, len(arr)):
                                    arr.pop()
                                arr.append(temp);
                                pnum = 1

                        else:
                            if c == ']':
                                temp = ''
                                for j in range(nowpos + 1, len(arr)):
                                    temp = temp + arr[j]
                                temp = "(" + temp + "){0,1}";
                                for j in range(nowpos, len(arr)):
                                    arr.pop()
                                arr.append(temp);
                            elif c == '}':
                                temp = ''
                                for j in range(nowpos + 1, len(arr)):
                                    temp = temp + arr[j]
                                temp = "(" + temp + ")";
                                for j in range(nowpos, len(arr)):
                                    arr.pop()
                                arr.append(temp);
                    else:
                        if c == ']':
                            temp = ''
                            for j in range(nowpos + 1, len(arr)):
                                temp = temp + arr[j]
                            temp = "(" + temp + "){0,1}";
                            for j in range(nowpos, len(arr)):
                                arr.pop()
                            arr.append(temp);
                        elif c == '}':
                            temp = ''
                            for j in range(nowpos + 1, len(arr)):
                                temp = temp + arr[j]
                            temp = "(" + temp + ")";
                            for j in range(nowpos, len(arr)):
                                arr.pop()
                            arr.append(temp);
                    expect.pop()
                    pos.pop()
                else:
                    return False
            elif c == '-':  # 操作判断
                if i + 1 < length:
                    k = 0
                    for j in range(i + 1, length):
                        # print(j,length)
                        if s[j].isalpha():
                            c = c + s[j]
                        else:
                            break
                        k = k + 1
                    i = i + k
                    if (frontisor):
                        temp = arr[-1] + "(" + c + ")"
                        arr.pop()
                        arr.append(temp)
                    elif s[i + 1] in SYMBOLS_R:
                        arr.append(c)
                    else:
                        arr.append("(" + c + ")")
                        '''
                    if (s[i + 1] == '|'):
                        pnum = pnum + 1
                        temp = arr[-1] + '|'
                        frontisor=True
                        arr.pop()
                        arr.append(temp)
                        '''
                else:
                    return False
            elif c.isalpha():  # 参数判断
                j = i
                while (i+1<length and (s[i + 1].isalpha() or s[i+1]=='-')):
                    i = i + 1
                if (i > j):
                    parameter = s[j:i + 1]
                else:
                    parameter = c
                if (frontisor):
                    temp = arr[-1] + "(" + self.ParameterToReg[parameter] + ")"
                    arr.pop()
                    arr.append(temp)
                else:
                    arr.append("(" + self.ParameterToReg[parameter] + ")")
                # print(s[i+1])
                '''
                if(i+1<length and s[i+1]=='|'):
                    pnum=pnum+1
                    frontisor=True
                    temp=arr[-1]+'|'
                    arr.pop()
                    arr.append(temp)
                '''
            elif c == "&":

                j = i + 1
                while (s[i] != '>'):
                    i = i + 1
                temp = s[j + 1:i]
                (begin, end) = temp.split('-')
                pwithiter = arr[-1] + "{" + begin + "," + end + "}"
                arr.pop()
                arr.append(pwithiter)
            elif c == "|":
                frontisor = True
                temp = ''
                for j in range(nowpos + 1, len(arr)):
                    temp = temp + arr[j]
                for j in range(nowpos + 1, len(arr)):
                    arr.pop()
                temp = temp + '|'
                arr.append(temp);
                pnum = pnum + 1
            elif c==' ':
                temp=arr[-1]+' '
                arr.pop()
                arr.append(temp)

            i = i + 1

        REG = ''
        for k in arr:
            REG = REG + k
        self.REG = REG
        if(len(expect)!=0):
            return False
        return True

#test=Templates('CommandTemplate.docx')